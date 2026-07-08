"""Validation framework: run after each pipeline execution, produces a
report the COO could read before approving a refresh.

Outputs three files: latest_report.json (machine-readable), latest_report.md
and latest_report.docx (both human-readable, same content -- a plain-English
summary up top for a non-technical reader, full per-check technical detail
below for engineering/audit review).

Usage:
    python -m validation.run
"""
import datetime as dt
import json

from pipeline.config import LOG_DIR, REPO_ROOT
from validation.checks import run_all_checks

REPORT_MD = REPO_ROOT / "validation" / "latest_report.md"
REPORT_DOCX = REPO_ROOT / "validation" / "latest_report.docx"
REPORT_JSON = REPO_ROOT / "validation" / "latest_report.json"

SOURCE_LABELS = {
    "pcc_residents": "PCC",
    "pcc_incidents": "PCC",
    "pcc_care_history": "PCC",
    "yardi_units": "Yardi",
    "yardi_leases": "Yardi",
    "adp_shifts": "ADP",
    "gbp_reviews": "Google Business Profile",
    "hubspot_leads": "HubSpot",
}

WHAT_WAS_CHECKED = [
    "Row count reconciliation from the raw source (Bronze layer) to the final output "
    "(Gold layer).",
    "Total revenue, total shift hours, and total resident-days match exactly between the "
    "raw exports and what's on the dashboard.",
    "Business rule checks:No overlapping leases for the same resident, no negative occupancy, no discharge before admits, "
    "no future dataed events, acuity scores within range."
     "Every reference between tables resolves correctly (e.g. every incident, lease, and "
    "shift points to a real resident/community that actually exists).",
]


def _plural(n: int, singular: str, plural: str | None = None) -> str:
    return singular if n == 1 else (plural or f"{singular}s")


def _anomaly_plain_english(a: dict) -> str:
    """Translates one raw anomaly record into a sentence a non-technical
    reader can act on (or confirm needs no action from them). Anomalies are
    now broken out per source file (see pipeline/run.py), so each sentence
    ends with exactly which monthly export it traces back to -- lets a
    reader (or the walkthrough) point at the one file responsible instead of
    just an aggregate table-wide count."""
    table, reason, n = a["table"], a["reason"], a["rows_affected"]
    source = SOURCE_LABELS.get(table, table)

    if reason == "acuity_score_out_of_range":
        message = (
            f"{n} {_plural(n, 'resident')} had an invalid acuity score in {source} (outside "
            "the 1-10 scale) -- the score was set aside pending correction; the rest of the "
            "resident's data is unaffected."
        )
    elif reason == "future_dated_discharge":
        message = (
            f"{n} {_plural(n, 'resident')} had an impossible future discharge date in "
            f"{source} -- set aside pending correction."
        )
    elif reason == "unresolvable_hourly_rate":
        message = (
            f"{n} shift {_plural(n, 'record')} in {source} had a pay rate that couldn't be "
            "determined -- excluded from labor cost totals pending correction."
        )
    elif reason == "duplicate_lead_id_conflicting_data":
        message = (
            f"{n} {_plural(n, 'lead')} in {source} shared a duplicate ID with conflicting "
            f"data -- set aside pending manual review with {source}'s data team."
        )
    elif reason == "unknown_community_id":
        message = (
            f"{n} {_plural(n, 'record')} in {source} {_plural(n, 'was', 'were')} tagged to a "
            "community that doesn't exist in our 14-community list -- excluded from reporting."
        )
    else:
        # Fallback for any reason not covered above -- still a plain sentence, not a
        # crash, and points the reader at the technical section for the underlying detail.
        message = f"{n} {_plural(n, 'record')} in {source} flagged for '{reason}' -- see technical detail below."

    source_file = a.get("source_file")
    return f"{message} (source: {source_file})" if source_file else message


def _build_summary(report: dict) -> dict:
    return {
        "verdict_ok": report["checks_failed"] == 0,
        "anomaly_lines": [_anomaly_plain_english(a) for a in report["pipeline_anomalies"]],
        "failing_checks": [c["check"] for c in report["checks"] if c["status"] == "fail"],
    }


def load_pipeline_log() -> dict:
    log_path = LOG_DIR / "latest_run.json"
    if not log_path.exists():
        raise SystemExit("No pipeline run found. Run `python -m pipeline.run` first.")
    return json.loads(log_path.read_text())


def run() -> dict:
    pipeline_log = load_pipeline_log()
    checks = run_all_checks(pipeline_log["bronze"], pipeline_log["silver"])

    failures = [c for c in checks if c["status"] == "fail"]
    report = {
        "generated_at": dt.datetime.now(dt.timezone.utc).isoformat(),
        "pipeline_run_at": pipeline_log["run_at"],
        "checks_run": len(checks),
        "checks_passed": len(checks) - len(failures),
        "checks_failed": len(failures),
        "checks": checks,
        "pipeline_anomalies": pipeline_log["anomalies"],
    }

    REPORT_JSON.write_text(json.dumps(report, indent=2, default=str))
    _write_markdown(report)
    _write_docx(report)
    return report


def _write_markdown(report: dict) -> None:
    summary = _build_summary(report)
    lines = [
        "# Pinewood Data Refresh — Validation Summary",
        "",
        f"Generated {report['generated_at']} UTC · based on the pipeline run at {report['pipeline_run_at']} UTC",
        "",
    ]

    if summary["verdict_ok"]:
        lines.append(f"### ✅ Safe to approve — all {report['checks_run']} accuracy and completeness checks passed.")
    else:
        lines.append(f"### ⚠️ NOT safe to approve without review — {report['checks_failed']} of {report['checks_run']} checks failed.")
        lines += ["", "Failing checks:"]
        lines += [f"- `{name}`" for name in summary["failing_checks"]]

    lines += ["", "**What was checked:**", ""]
    lines += [f"- {b}" for b in WHAT_WAS_CHECKED]

    lines.append("")
    if summary["anomaly_lines"]:
        n = len(summary["anomaly_lines"])
        lines.append(
            f"**{n} data-quality {'issue' if n == 1 else 'issues'} found in this month's source "
            "exports, handled automatically -- none require action from you:**"
        )
        lines.append("")
        lines += [f"- {line}" for line in summary["anomaly_lines"]]
    else:
        lines.append("**No data-quality issues found in this month's source exports.**")

    lines += ["", "---", "", "## Full technical detail (for engineering / audit review)", ""]
    lines += ["| Check | Status | Severity | Recommended action | Detail |", "|---|---|---|---|---|"]
    for c in report["checks"]:
        detail = json.dumps(c["detail"], default=str)
        if len(detail) > 120:
            detail = detail[:117] + "..."
        lines.append(f"| {c['check']} | {c['status'].upper()} | {c['severity']} | {c['action']} | `{detail}` |")

    lines += ["", "### Anomalies detected during ingestion (raw log)", ""]
    if report["pipeline_anomalies"]:
        lines.append("| Table | Reason | Source file | Rows affected | Severity | Action |")
        lines.append("|---|---|---|---|---|---|")
        for a in report["pipeline_anomalies"]:
            lines.append(
                f"| {a['table']} | {a['reason']} | {a.get('source_file', 'n/a')} | "
                f"{a['rows_affected']} | {a['severity']} | {a['action']} |"
            )
    else:
        lines.append("None.")

    REPORT_MD.write_text("\n".join(lines) + "\n")


def _write_docx(report: dict) -> None:
    from docx import Document
    from docx.shared import RGBColor

    summary = _build_summary(report)
    doc = Document()

    doc.add_heading("Pinewood Data Refresh — Validation Summary", level=1)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Generated {report['generated_at']} UTC · based on the pipeline run at {report['pipeline_run_at']} UTC"
    ).italic = True

    verdict = doc.add_paragraph()
    if summary["verdict_ok"]:
        run = verdict.add_run(f"Safe to approve — all {report['checks_run']} accuracy and completeness checks passed.")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x7F, 0x37)
    else:
        run = verdict.add_run(
            f"NOT safe to approve without review — {report['checks_failed']} of {report['checks_run']} checks failed."
        )
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        doc.add_paragraph("Failing checks:")
        for name in summary["failing_checks"]:
            doc.add_paragraph(name, style="List Bullet")

    doc.add_heading("What was checked", level=2)
    for b in WHAT_WAS_CHECKED:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Data-quality issues found", level=2)
    if summary["anomaly_lines"]:
        n = len(summary["anomaly_lines"])
        doc.add_paragraph(
            f"{n} data-quality {'issue' if n == 1 else 'issues'} found in this month's source "
            "exports, handled automatically -- none require action from you:"
        )
        for line in summary["anomaly_lines"]:
            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("No data-quality issues found in this month's source exports.")

    doc.add_heading("Full technical detail (for engineering / audit review)", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, header in enumerate(["Check", "Status", "Severity", "Action", "Detail"]):
        table.rows[0].cells[i].text = header
    for c in report["checks"]:
        detail = json.dumps(c["detail"], default=str)
        if len(detail) > 200:
            detail = detail[:197] + "..."
        row = table.add_row().cells
        row[0].text = c["check"]
        row[1].text = c["status"].upper()
        row[2].text = c["severity"]
        row[3].text = c["action"]
        row[4].text = detail

    doc.add_heading("Anomalies detected during ingestion (raw log)", level=2)
    if report["pipeline_anomalies"]:
        t2 = doc.add_table(rows=1, cols=6)
        t2.style = "Light Grid Accent 1"
        for i, header in enumerate(["Table", "Reason", "Source file", "Rows affected", "Severity", "Action"]):
            t2.rows[0].cells[i].text = header
        for a in report["pipeline_anomalies"]:
            row = t2.add_row().cells
            row[0].text = a["table"]
            row[1].text = a["reason"]
            row[2].text = a.get("source_file", "n/a")
            row[3].text = str(a["rows_affected"])
            row[4].text = a["severity"]
            row[5].text = a["action"]
    else:
        doc.add_paragraph("None.")

    doc.save(str(REPORT_DOCX))


if __name__ == "__main__":
    result = run()
    print(f"{result['checks_passed']}/{result['checks_run']} checks passed.")
    if result["checks_failed"]:
        print(f"{result['checks_failed']} FAILED — see validation/latest_report.md / .docx")
    else:
        print("See validation/latest_report.md (or .docx) for the full report.")

